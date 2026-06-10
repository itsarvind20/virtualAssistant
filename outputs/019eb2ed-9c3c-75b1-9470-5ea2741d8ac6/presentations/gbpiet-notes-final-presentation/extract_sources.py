import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


DOCX_PATH = Path(r"C:\Users\HP\Desktop\GBPIET_NOTES-main\GBPIET_NOTES-main\GBPIET_Notes_Final_Project_Report.docx")
PDF_PATH = Path(r"C:\Users\HP\Desktop\EDUCONNECT newwwwww_progressPPT.pdf")
WORKSPACE = Path(r"C:\Users\HP\Documents\GitHub\virtualAssistant\outputs\019eb2ed-9c3c-75b1-9470-5ea2741d8ac6\presentations\gbpiet-notes-final-presentation")


def clean(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text


def extract_docx():
    doc = Document(DOCX_PATH)
    paragraphs = []
    headings = []
    for p in doc.paragraphs:
        text = clean(p.text)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        item = {"style": style, "text": text}
        paragraphs.append(item)
        if "Heading" in style or text.isupper() or re.match(r"^\d+(\.\d+)*\s+", text):
            headings.append(item)

    tables = []
    for ti, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"index": ti, "rows": rows})

    return {"paragraphs": paragraphs, "headings": headings, "tables": tables}


def extract_pdf():
    reader = PdfReader(str(PDF_PATH))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        text = clean(page.extract_text() or "")
        pages.append({"page": idx, "text": text})
    return {"pages": pages}


def main():
    docx = extract_docx()
    pdf = extract_pdf()
    data = {"docx": docx, "pdf": pdf}
    (WORKSPACE / "source-extract.json").write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

    lines = ["# DOCX HEADINGS"]
    lines.extend(f"- {h['text']} [{h['style']}]" for h in docx["headings"][:120])
    lines.append("\n# DOCX BODY SAMPLE")
    lines.extend(f"- {p['text']}" for p in docx["paragraphs"][:220])
    lines.append("\n# DOCX TABLES")
    for table in docx["tables"]:
        lines.append(f"Table {table['index']}")
        for row in table["rows"][:12]:
            lines.append(" | ".join(row))
        lines.append("")
    lines.append("\n# PDF PAGES")
    for page in pdf["pages"]:
        lines.append(f"Page {page['page']}: {page['text'][:1200]}")
    (WORKSPACE / "source-digest.txt").write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps({
        "docx_paragraphs": len(docx["paragraphs"]),
        "docx_headings": len(docx["headings"]),
        "docx_tables": len(docx["tables"]),
        "pdf_pages": len(pdf["pages"]),
    }, indent=2))


if __name__ == "__main__":
    main()
